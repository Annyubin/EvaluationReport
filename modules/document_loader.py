"""
HWP 및 DOCX 문서 로더 모듈
"""

import logging
import os
import tempfile
from typing import Optional, List
from pathlib import Path
import subprocess
import sys

# 로깅 설정
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """HWP 및 DOCX 문서 로더 클래스"""
    
    def __init__(self):
        self.supported_formats = ["hwp", "docx"]
    
    def get_supported_formats(self) -> List[str]:
        """지원하는 파일 형식 반환"""
        return self.supported_formats.copy()
    
    def load_document(self, file_path: str) -> Optional[str]:
        """문서 파일을 로드하여 텍스트 추출"""
        try:
            path_obj = Path(file_path)
            if not path_obj.exists():
                logger.error(f"파일이 존재하지 않습니다: {file_path}")
                return None
            
            file_extension = path_obj.suffix.lower()
            
            if file_extension == '.hwp':
                return self._extract_hwp_text(path_obj)
            elif file_extension == '.docx':
                return self._extract_docx_text(path_obj)
            else:
                logger.error(f"지원하지 않는 파일 형식입니다: {file_extension}")
                return None
                
        except Exception as e:
            logger.error(f"문서 로드 중 오류 발생: {e}")
            return None
    
    def _extract_hwp_text(self, file_path: Path) -> Optional[str]:
        """HWP 파일에서 텍스트 추출"""
        try:
            # hwp5txt 명령어 사용
            result = subprocess.run(
                ['hwp5txt', str(file_path)],
                capture_output=True,
                text=True,
                encoding='utf-8'
            )
            
            if result.returncode == 0:
                return result.stdout
            else:
                logger.warning(f"hwp5txt 실패, 대체 방법 시도: {result.stderr}")
                return self._extract_hwp_text_fallback(file_path)
                
        except FileNotFoundError:
            logger.warning("hwp5txt가 설치되지 않았습니다. 대체 방법을 사용합니다.")
            return self._extract_hwp_text_fallback(file_path)
        except Exception as e:
            logger.error(f"HWP 텍스트 추출 중 오류: {e}")
            return None
    
    def _extract_hwp_text_fallback(self, file_path: Path) -> Optional[str]:
        """HWP 텍스트 추출 대체 방법"""
        try:
            # Python HWP 라이브러리 사용 (설치된 경우)
            import hwp5
            from hwp5.xmlmodel import Hwp5File
            
            hwp_file = Hwp5File(str(file_path))
            text_content = []
            
            for paragraph in hwp_file.bodytext.sections[0].paragraphs:
                for text in paragraph.texts:
                    if hasattr(text, 'content'):
                        text_content.append(text.content)
            
            return '\n'.join(text_content)
            
        except ImportError:
            logger.error("HWP 처리를 위한 라이브러리가 설치되지 않았습니다.")
            return None
        except Exception as e:
            logger.error(f"HWP 대체 방법 실패: {e}")
            return None
    
    def _extract_docx_text(self, file_path: Path) -> Optional[str]:
        """DOCX 파일에서 텍스트 추출"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # 테이블에서도 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            return '\n'.join(text_content)
            
        except ImportError:
            logger.error("python-docx 라이브러리가 설치되지 않았습니다.")
            return None
        except Exception as e:
            logger.error(f"DOCX 텍스트 추출 중 오류: {e}")
            return None
    
    def validate_file_format(self, file_path: str) -> bool:
        """파일 형식 검증"""
        try:
            file_extension = Path(file_path).suffix.lower()
            return file_extension in ['.hwp', '.docx']
        except Exception:
            return False
    
    def get_file_info(self, file_path: str) -> dict:
        """파일 정보 반환"""
        try:
            path_obj = Path(file_path)
            if not path_obj.exists():
                return {}
            
            return {
                'name': path_obj.name,
                'size': path_obj.stat().st_size,
                'extension': path_obj.suffix.lower(),
                'is_supported': self.validate_file_format(str(path_obj))
            }
        except Exception as e:
            logger.error(f"파일 정보 조회 중 오류: {e}")
            return {} 